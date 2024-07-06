import pandas as pd
from docx import Document
from docx.shared import Inches , Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from pillow_heif import register_heif_opener

register_heif_opener()
df = pd.read_excel('fnfin.xlsx')  
# df=pd.read_ex('fix.csv')
doc = Document()

# print(df.columns)

table = doc.add_table(rows=0, cols=3)
for row in table.rows:
    row.height_rule = Inches(5) 
df['Submit yearbook quote to be printed under your image (feel free to show your creativity!):\n\n(ensure that it is under 70 characters and does not contain any profanity/vulgar language)'].fillna('Lorem ipsum dolor sit amet, consectetur adipiscing elit', inplace=True)
df['Submit picture of yourself (please ensure it is decent, clear and well lit, preferably 3:4 image ratio, or passport sized) you can only upload the picture once as editing photos is not allowed on google forms.'].fillna('https://drive.google.com/file/d/1dxSlFP84x_jJbAH2SJMjOrwW_S87b6iy/view?usp=sharing', inplace=True)

ratio = 1
# quotes = 'Lorem ipsum dolor sit amet, consectetur adipiscing elit'
for i in range(0, len(df), 3):
    row_cells = table.add_row().cells
    for j in range(3):
        if i + j < len(df):
            # print(i,j)
            photo_link = df.loc[i + j, 'Submit picture of yourself (please ensure it is decent, clear and well lit, preferably 3:4 image ratio, or passport sized) you can only upload the picture once as editing photos is not allowed on google forms.']
            name = df.loc[i + j, 'Name']
            surname = df.loc[i + j, 'BITS ID (this form is only for students enrolled in the year 2020)']

                
            quotes = df.loc[i + j, 'Submit yearbook quote to be printed under your image (feel free to show your creativity!):\n\n(ensure that it is under 70 characters and does not contain any profanity/vulgar language)']

            
            
        
            paragraph = row_cells[j].add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            if(photo_link == 'https://drive.google.com/file/d/1dxSlFP84x_jJbAH2SJMjOrwW_S87b6iy/view?usp=sharing'):
                import cv2 
                image_data = cv2.imread('Default.jpg')
                img = Image.fromarray(image_data)   
                
                desired_aspect_ratio = 1/ratio
                img_aspect_ratio = img.width / img.height

                if img_aspect_ratio > desired_aspect_ratio:
                    new_width = int(img.height * desired_aspect_ratio)
                    left_margin = (img.width - new_width) // 2
                    right_margin = img.width - new_width - left_margin
                    crop_box = (left_margin, 0, img.width - right_margin, img.height)
                else:
                    new_height = int(img.width / desired_aspect_ratio)
                    top_margin = (img.height - new_height) // 4
                    bottom_margin = img.height - new_height - top_margin*3
                    crop_box = (0, top_margin, img.width, img.height - bottom_margin)

                cropped_img = img.crop(crop_box)
                new_width = img.width
                new_height = int(new_width * ratio)

                img = cropped_img.resize((new_width, new_height))


                img_io = BytesIO()
                img.save(img_io, format='JPEG')
                img_io.seek(0) 
                run.add_picture(img_io, width=Inches(1.5))            
                paragraph.add_run('\n')
                paragraph.add_run(f'{surname}\n').italic = True
                paragraph.add_run(f'{name}\n').bold = True
                quotes_run = paragraph.add_run(quotes)
                quotes_run.font.size = Pt(10)
                continue
            import requests
            from io import BytesIO
            
            file_id = ''
            try:
                file_id = photo_link.split('=')[1]
                download_url = f'https://drive.google.com/uc?export=download&id={file_id}'
                

                # print(download_url)
                response = requests.get(download_url)
                image_data = BytesIO(response.content)
                from PIL import Image
                img = Image.open(image_data)
                
                desired_aspect_ratio = 1/ratio
                img_aspect_ratio = img.width / img.height

                if img_aspect_ratio > desired_aspect_ratio:
                    new_width = int(img.height * desired_aspect_ratio)
                    left_margin = (img.width - new_width) // 2
                    right_margin = img.width - new_width - left_margin
                    crop_box = (left_margin, 0, img.width - right_margin, img.height)
                else:
                    new_height = int(img.width / desired_aspect_ratio)
                    top_margin = (img.height - new_height) // 4
                    bottom_margin = img.height - new_height - top_margin*3
                    crop_box = (0, top_margin, img.width, img.height - bottom_margin)

                cropped_img = img.crop(crop_box)
                new_width = img.width
                new_height = int(new_width * ratio)

                img = cropped_img.resize((new_width, new_height))
                img = img.convert('RGB')
                # print(img)
                img_io = BytesIO()
                try:
                    img.save(img_io,format='JPEG')
                except:
                    try:
                        img.save(img_io,format='PNG')
                    except:
                        img.save(img_io,format='HEIC')
                img_io.seek(0)
                run.add_picture(img_io, width=Inches(1.5))            
                paragraph.add_run('\n')
                paragraph.add_run(f'{surname}\n').italic = True
                paragraph.add_run(f'{name}\n').bold = True
                quotes_run = paragraph.add_run(quotes)
                quotes_run.font.size = Pt(10)
            except:   
                # print(name, surname, photo_link, quotes)
                run.add_picture('Default.jpg', width=Inches(1.5))            
                paragraph.add_run('\n')
                paragraph.add_run(f'{surname}\n').italic = True
                paragraph.add_run(f'{name}\n').bold = True
                quotes_run = paragraph.add_run(quotes)
                quotes_run.font.size = Pt(10)
                print(name, surname, photo_link, quotes)
            
        # if((i-3)%6 == 0):
        #     doc.add_page_break()
              
                   
# 1/0.95
doc.save('fnFin1.docx')